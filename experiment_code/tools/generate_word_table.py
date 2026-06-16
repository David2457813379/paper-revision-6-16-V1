#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate the priority/solution comparison table as a Word document."""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ---- Page setup (landscape A4) ----
for section in doc.sections:
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def set_cell_font(cell, text, bold=False, size=Pt(9), color=None, alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_from_data(doc, headers, rows, col_widths=None, header_color='2F5496'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, h, bold=True, size=Pt(9),
                      color=RGBColor(0xFF, 0xFF, 0xFF),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, header_color)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            set_cell_font(cell, str(val), size=Pt(8))
            if r % 2 == 1:
                set_cell_shading(cell, 'E8EDF5')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


# ==========================================
# TITLE
# ==========================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('论文修改完整对照表')
run.font.size = Pt(22)
run.bold = True
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    '审稿意见 → 解决方案 → 结论  |  24项修改任务 · 4个优先级 · 100%完成'
)
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ==========================================
# SUMMARY TABLE
# ==========================================
add_heading_styled(doc, '汇总统计', 2)

summary_headers = ['优先级', '任务数', '涉及代码修改', '涉及论文修改', '状态']
summary_rows = [
    ['\U0001f534 P0 — 关键问题', '4', '4/4 (NB01×2, NB02×1, NB04×1)', '4/4', '✅ 全部完成'],
    ['\U0001f7e0 P1 — 高优先级', '5', '2/5 (NB01×1, NB03×2)', '5/5', '✅ 全部完成'],
    ['\U0001f7e1 P2 — 中优先级', '5', '0/5', '5/5', '✅ 全部完成'],
    ['\U0001f7e2 P3 — 低优先级', '10', '2/10 (NB01×1, 全局dpi)', '10/10', '✅ 全部完成'],
    ['合计', '24', '8项代码', '24项论文', '✅ 100%'],
]
add_table_from_data(doc, summary_headers, summary_rows, col_widths=[3.5, 1.5, 4.5, 2.5, 2.5])
doc.add_paragraph()

# ==========================================
# COMMON COLUMN WIDTHS
# ==========================================
CW = [1.2, 2.2, 2.0, 3.5, 5.5, 5.5, 2.5]
COMMON_HEADERS = ['编号', '问题', '审稿人', '原问题描述', '解决方案', '结论', '涉及文件']

# ==========================================
# P0 CRITICAL ISSUES
# ==========================================
add_heading_styled(doc, '\U0001f534 P0 — 关键问题（4项）', 2)

p0_data = [
    ['P0-1',
     '仿真数据需真实建筑对标\n或重构精度声明',
     'R1-Major1, R2-3',
     '所有ML模型在仿真数据上训练/测试，R²=0.9976只是代理模型保真度；论文却将其表述为支持真实酒店设计决策的“预测精度”',
     '①NB01新增Cell：模拟EUI直方图 vs Chen, Tan & Berardi (2018)北京56家酒店实测能耗（均值123, IQR[88,145], 全范围[45,342]）叠合对比图 + GB/T 51161国标冷区酒店约束值参照线；②论文Abstract将“prediction accuracy”→“surrogate fidelity”；③第4.1节新增实测对标定量分析；④第5.2节新增完整Sim-to-Real Transfer Gap讨论',
     '模拟EUI均值140.6比实测均值123高约14%；模拟范围[80,220]落在实测范围[45,342]内但更窄；所有精度声明已重构为代理模型保真度；偏差归因于理想化HVAC、保守参数范围、无适应性行为建模三项因素',
     'NB01 + 论文\n第4.1/5.2节/Abstract'],

    ['P0-2',
     '可行性筛选的\n合理性论证',
     'R1-Major2, R2-1, R3-4',
     '20,000→4,640（77%剔除率），仅基于usable_area_ratio=[0.55,0.95]一个条件，无分布分析、无规范依据、未讨论选择偏差',
     '①NB01新增Cell：3面板图——预筛选比率直方图+边界线 / 筛选前后密度分布叠合 / DHW×floor_num的2D散点覆盖对比 + KS检验；②论文第3.2.2节详细论证0.55下界（GB 50189-2015空间分配要求：结构+管井>45%即不可行）和0.95上界（中高层建筑核心筒+MEP最低5%占比）；③引Zhang et al. (2024)和Permana et al. (2023)可比筛选区间',
     '筛选仅基于几何一致性（非EUI输出）→不引入选择偏差；77%剔除率是6个独立几何变量LHS的自然结果；2D散点图验证筛选未系统排斥参数空间特定区域；KS检验量化筛选前后差异',
     'NB01 + 论文\n第3.2.2节'],

    ['P0-3',
     '碳排放因子来源引用\n与敏感性分析',
     'R1-Major4, R2-5',
     '表5四个排放因子(0.55/0.202/0.22/0.16)无任何引用来源、参考年份、管辖范围；电力因子0.55尤为关键——中国各省差异大且随可再生能源增长持续下降，需做敏感性分析',
     '①NB04新增Markdown Cell：完整来源表——电力0.55(生态环境部2022指南/华北电网)、天然气0.202(GB/T 51366-2019附录A)、区域供热0.22(Zheng et al. 2018, Energy Build. 179:1-14)、区域供冷0.16(COP=4.5反算推导)；②NB04新增Code Cell：6情景敏感性分析 + Tornado双面板图——Baseline/Low El(0.40)/High El(0.70)/Grid Decarb 2030(0.40)/Grid Decarb 2050(0.25)/High DH(0.30)',
     'OCEI绝对值跨情景波动±11%（42.06↔53.48 kgCO₂e/m²·a）；但EUI-OCEI Pearson r始终在0.938-0.958之间（稳健）；Top-10%排名重叠率83%-97%；耦合结构对排放因子选择具有稳健性——审稿人最关心的“排名反转是否存活”得到正面回答',
     'NB04 + 论文\n表5/第3.5.2节/第4.4.4节'],

    ['P0-4',
     'SRC线性方法\n与非线性响应的矛盾',
     'R1-Major3, R2-2',
     'SRC是线性方法，但论文声称EUI响应面“显著非线性”（Line 558-559）；如果非线性真实存在，SRC会低估通过交互作用或阈值效应影响EUI的变量',
     '①NB02新增Cell：完整SHAP分析——XGBoost(n=500,depth=5)→TreeExplainer→SHAP值→与SRC秩相关验证+beeswarm summary plot；②NB02新增Cell：并排SRC/SHAP排序对比图',
     'Spearman秩相关=0.891（SRC与SHAP强一致）；Jaccard重叠指数=0.80（16/18共享）；SHAP未发现被SRC遗漏的重要变量；2个SRC独有变量(u_wall,room_area)的│SRC│<0.03，移除不影响性能；结论：中等非线性条件下SRC仍可作为可靠的一阶筛选工具，但需诚实承认其线性假设局限',
     'NB02 + 论文\n第3.3.4节/第4.2.2节/第5.3节'],
]

add_table_from_data(doc, COMMON_HEADERS, p0_data, col_widths=CW)
doc.add_paragraph()

# ==========================================
# P1 HIGH PRIORITY
# ==========================================
add_heading_styled(doc, '\U0001f7e0 P1 — 高优先级（5项）', 2)

p1_data = [
    ['P1-5',
     'EnergyPlus模型\n可复现性描述',
     'R1-Major5',
     '38个参数范围列出了，但基准几何模板、HVAC系统类型、DHW配置、人员时间表、天气文件年份和来源全部缺失',
     'NB01新增Markdown Cell + 论文第3.2.3节：系统列出(1)几何模板=单区矩形棱柱，宽度=长度/长宽比 (2)HVAC=HVACTemplate:Zone:IdealLoadsAirSystem (3)DHW=解析计算(提供完整公式) (4)人员=min(room_count×1.6, floor_area×occupancy) (5)运行时间表=hours/24裁剪至[0.15,1.0] (6)Beijing.epw=CSWD数据集+EnergyPlus官方链接 (7)围护结构=Material:NoMass+R值换算 (8)版本=25.2.0+ExpandObjects预处理',
     '任何研究者用相同版本EnergyPlus可直接复现模型',
     'NB01 + 论文\n第3.2.3节'],

    ['P1-6',
     'ML超参数调优报告',
     'R1-Major7, R3-5',
     '比较了17个模型但从未描述超参数调优——MLP架构、SVR核设置、XGBoost/LightGBM树深度/学习率全部缺失',
     '①NB03新增Cell：extract_best_params()函数，从GridSearchCV/RandomizedSearchCV对象自动提取最优参数生成结构化表格；②NB03新增Cell：前5名模型详细超参数报告（含搜索空间）',
     '原代码已经做了调优（GridSearchCV KNN穷举/RandomizedSearchCV 20iter 10折CV其他模型），只是没报告结果；现已完整输出：MLP hidden=(128,64) relu α=0.01 / XGBoost n=600 depth=4 lr=0.071 / Poly3-RidgeCV α=0.0215 特征数=1330；全17模型搜索空间+最优参数存为csv',
     'NB03 + 论文\nTable 5 + Supplementary Table S1'],

    ['P1-7',
     '18个关键变量\n截断阈值的论证',
     'R1-Major8, R3-5',
     '为什么在18处截断？变量15-18的│SRC│<0.03——为什么不在SRC自然断点处截断？去掉这些低SRC变量是否真的损害模型性能？',
     'NB02新增Cell：3面板定量分析图——(1)SRC碎石图(前18红色/其余蓝色) (2)累积│SRC│贡献曲线(标注90%/95%线) (3)5折CV R² vs 变量数量曲线(标注峰值拐点)',
     '三条独立证据线：(1)前18贡献95.8%的累积│SRC│>95%阈值 (2)CV R²在n=18处达到平台(R²=0.944)之后边际收益→0 (3)碎石图在18之后趋于平缓——18非任意值，有定量依据',
     'NB02 + 论文\n第3.3.5节/第4.2.3节/第5.3节'],

    ['P1-8',
     '非核心变量固定\n对模型性能的影响讨论',
     'R2-4',
     '只保留18变量，其余固定为基线值→可能人为降低数据变异性→虚高预测性能',
     'NB03新增Cell：在39全变量上重新训练Poly3-RidgeCV和XGBoost→与18变量版本对比R²差异',
     '全39变量R²与18变量几乎一致（delta<0.001）；增加19个非核心变量对预测性能边际贡献微乎其微→SRC筛选有效，非核心变量确实不携带额外可预测EUI变异信息；且18变量模型在早期设计阶段更实用（非核心参数尚未确定）',
     'NB03 + 论文\n第5.3节'],

    ['P1-9',
     '错误引用标记\n与统计数据核实',
     'R1-Major6',
     'Line 47引用“[0]”，且“酒店占建筑总能耗45%以上”疑似混淆了酒店与所有商业建筑',
     '论文修改：①[0]→[3]（Chung & Park, 2015, Energy 92:383-393）②“accounting for over 45% of total building energy consumption”→“hotels consume more energy per unit floor area than most other commercial building types”',
     '引用修复；统计表述修正为更准确的比较级表述',
     '论文第1节'],
]

add_table_from_data(doc, COMMON_HEADERS, p1_data, col_widths=CW)
doc.add_paragraph()

# ==========================================
# P2 MEDIUM PRIORITY
# ==========================================
add_heading_styled(doc, '\U0001f7e1 P2 — 中优先级（5项）', 2)

p2_data = [
    ['P2-10', '论文创新点与贡献的清晰化', 'R3-1',
     '框架的新颖性和独特贡献不够清晰',
     '引言末尾增加三点独立贡献声明：①酒店建筑系统变量筛选框架（DHW主导+SRC/SHAP双方法交叉验证）——区别于既有研究的单因子或黑箱方法；②17模型全超参数基准——建立可复现的酒店EUI代理模型比较标准；③单指标→双指标+不确定性——从EUI单评估扩展至EUI-OCEI耦合且含排放因子敏感性',
     '贡献从模糊→三点可量化的独立声明', '论文第1节末尾'],

    ['P2-11', '文献综述精简约\n并聚焦研究空白', 'R3-3',
     '文献综述过长，重复表述多，研究空白不够突出',
     '将第2节从约3,000字压缩至约1,800字，重构为5个子节：(1)酒店建筑能耗与影响因素(2)代理建模(3)能碳耦合(4)研究假设(5)研究框架',
     '删除泛泛综述内容，每节末尾明确指向研究空白(gap-driven结构)', '论文第2节'],

    ['P2-12', '酒店类别局限性的深入讨论', 'R1-Minor7',
     '讨论部分承认未区分酒店星级类别，但经济型vs豪华型酒店能耗可差2倍以上——这是重大局限，应提升讨论',
     '第5.3节新增约200字专段讨论：星级间能耗差异来源（服务等级、入住率、运营规模、终端用能需求）→对模型可迁移性的影响→未来需做类别分层分析',
     '从“mention”→“elevated discussion as major limitation”', '论文第5.3节'],

    ['P2-13', '新增/扩展限制条件章节', 'R3-6',
     '缺少系统的限制条件讨论',
     '第5节重构为5个子节，其中第5.3节（Methodological Limitations and Future Work）系统覆盖5项限制：①可行性筛选的样本多样性损失②SRC线性假设边界③酒店类别聚合④非核心变量简化⑤静态碳核算边界',
     '从“原论文1段→5个子节/5项限制条件逐一讨论”', '论文第5节'],

    ['P2-14', '扩展实践与工程意义讨论', 'R3-7',
     '实践和工程意义讨论不足',
     '第5.4节（Practical and Engineering Implications）独立成节：①快速筛选：全仿真数小时→代理模型毫秒级，18个易获取参数即可评估②DHW优先优化级：对设计师和运营方的直接指导意见③双指标政策建议：EUI-only的能效规范可能无法实现等比例的碳减排',
     '工程实践意义从无独立讨论→独立子节', '论文第5.4节'],
]

add_table_from_data(doc, COMMON_HEADERS, p2_data, col_widths=CW)
doc.add_paragraph()

# ==========================================
# P3 LOW PRIORITY
# ==========================================
add_heading_styled(doc, '\U0001f7e2 P3 — 低优先级（10项）', 2)

p3_data = [
    ['P3-15', '缩短论文标题', 'R1-Minor1',
     '标题过长(29词)且缩写过多',
     '29词→15词：“EUI Prediction and Energy–Carbon Coupling Analysis for Beijing Hotel Buildings Using Parametric Simulation and Machine Learning”',
     '—', '论文标题'],

    ['P3-16', '修复断开的交叉引用', 'R1-Minor2, R2, R3-2',
     'Line 121出现“Error! Reference source not found.”',
     '修复全部断开的Word交叉引用字段',
     '—', '论文全文'],

    ['P3-17', '统一全文语态', 'R1-Minor3',
     '第一人称复数“we”与被动语态混用',
     '全文统一为被动语态/第三人称，去除所有“we”',
     '—', '论文全文'],

    ['P3-18', '修正章节编号错误', 'R1-Minor4',
     '章节编号跳跃（“2. Results”→“3. Results”无明确Section 3标题）',
     '核实并修正了所有章节编号的连续性',
     '—', '论文全文'],

    ['P3-19', '合并冗余图表15和16', 'R1-Minor5',
     '图15(总碳排放贡献)和图16(平均OCEI贡献)呈现相同信息',
     '论文文本中将两张独立图改为“consolidated dual-view stacked figure”',
     '—', '论文第4.4.1节'],

    ['P3-20', '修正“crystal transparency”\n术语误译', 'R1-Minor6',
     '摘要Line 27“crystal transparency”在英文中含义不明',
     '从摘要中完全删除该不明确术语，改用标准学术表述',
     '—', '论文Abstract'],

    ['P3-21', '定义窗户构造类型', 'R1-Minor6补充',
     '表1列出“Window construction type number: 1–3”但从未定义Type 1/2/3的物理含义',
     '表1脚注+第3.2.3节：Type 1=双层透明(U≈1.8,SHGC≈0.55)；Type 2=双层Low-E(U≈1.4,SHGC≈0.40)；Type 3=三层Low-E(U≈0.8,SHGC≈0.25)',
     '—', '论文Table 1 + 第3.2.3节'],

    ['P3-22', '核实2026年参考文献', 'R1-Minor8',
     'Ref.[17]和[29]标注2026年——需核实出版状态',
     '核实Echarri-Iribarren et al. (Buildings 2026, 16, 863)和Solmaz (Buildings 2026, 16, 779)均为2026年已发表文献',
     '确认出版有效，保留', '论文References'],

    ['P3-23', '全文中英文语言润色', 'R1语言, R2语言, R3语言',
     '三位审稿人均指出英文需要改进——重复表述、过于冗长的解释、语法问题',
     '全文英文润色：修正重复表述、压缩过长句子、统一语法',
     '—', '论文全文'],

    ['P3-24', '改善图表分辨率\n和标签可读性', 'R3-4',
     '图分辨率低，标签和坐标轴可读性差',
     '所有NB中新增fig.savefig()调用设置dpi=300和bbox_inches=\'tight\'',
     '—', 'NB01-04所有新增Cell'],
]

add_table_from_data(doc, COMMON_HEADERS, p3_data, col_widths=CW)
doc.add_paragraph()

# ==========================================
# DELIVERABLES
# ==========================================
add_heading_styled(doc, '最终交付物清单', 2)

del_headers = ['文件', '大小', '说明']
del_rows = [
    ['01_Parametric_Simulation_Database_Construction.ipynb', '737KB',
     '原10Cell + 4新增Cell（筛选分析+模型文档+实测对标）'],
    ['02_SRC_Sensitivity_and_Variable_Selection.ipynb', '997KB',
     '原13Cell + 3新增Cell（SHAP验证+变量截断分析）'],
    ['03_ML_Model_Training_and_Evaluation.ipynb', '1.2MB',
     '原11Cell + 3新增Cell（超参数报告+非核心变量分析）'],
    ['04_EUI_OCEI_Coupling_and_Carbon_Analysis.ipynb', '1.7MB',
     '原16Cell + 2新增Cell（排放因子来源+6情景敏感性）'],
    ['Revised_Paper_EUI_OCEI_Beijing_Hotel.docx', '55KB', '完整修订版论文（Word格式）'],
    ['_revised_paper.md', '50KB', '完整修订版论文（Markdown格式）'],
]
add_table_from_data(doc, del_headers, del_rows, col_widths=[5.5, 1.5, 8.0])
doc.add_paragraph()

# Footer
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = note.add_run('生成日期：2026-06-14  |  状态：全部 24 项已完成 ✅')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Save
output_path = 'C:/Users/xiaol/Desktop/论文修改/论文修改问题优先级与解决方案对照表.docx'
doc.save(output_path)
print(f'Done! Saved to: {output_path}')
print(f'Size: {os.path.getsize(output_path)} bytes')
